from flask import Blueprint, render_template, session, jsonify, request, send_file, current_app, make_response
import os
from werkzeug.utils import secure_filename
from sds_translator_v4 import SDSTranslator
from database import get_db_path
from utils import AVAILABLE_LANGUAGES

main_bp = Blueprint('main', __name__)

@main_bp.route('/')
def index():
    """Main page with default template loaded."""
    default_template = "layout-placeholders-fixed-v2.html"
    template_loaded = False
    
    if os.path.exists(default_template):
        try:
            import shutil
            upload_folder = current_app.config['UPLOAD_FOLDER']
            session_id = os.urandom(16).hex()
            session['default_template_id'] = session_id
            
            default_path = os.path.join(upload_folder, f"default_{session_id}.html")
            shutil.copy(default_template, default_path)
            
            session['uploaded_file'] = default_path
            session['original_filename'] = default_template
            session['is_default'] = True
            template_loaded = True
        except Exception as e:
            print(f"Could not load default template: {e}")
    
    return render_template('index.html', 
                         languages=AVAILABLE_LANGUAGES, 
                         template_loaded=template_loaded, 
                         default_template=default_template)


@main_bp.route('/api/upload', methods=['POST'])
def upload_file():
    upload_folder = current_app.config['UPLOAD_FOLDER']
    
    # Validate file presence
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    # Validate filename
    if not file.filename or file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    # Security: Use secure_filename and validate extension
    filename = secure_filename(file.filename)
    
    # Validate file extension (whitelist approach)
    allowed_extensions = {'.html', '.htm'}
    file_ext = os.path.splitext(filename)[1].lower()
    if file_ext not in allowed_extensions:
        return jsonify({'error': 'Invalid file type. Only HTML files are allowed.'}), 400
    
    # Additional security: Check for path traversal attempts
    if '..' in filename or '/' in filename or '\\' in filename:
        return jsonify({'error': 'Invalid filename detected'}), 400
    
    # Validate file size (limit to 16MB as per Flask config)
    file.seek(0, 2)  # Seek to end
    file_size = file.tell()
    file.seek(0)  # Reset to beginning
    max_size = current_app.config.get('MAX_CONTENT_LENGTH', 16 * 1024 * 1024)
    if file_size > max_size:
        return jsonify({'error': f'File too large. Maximum size is {max_size // (1024*1024)}MB'}), 400
    
    if file and filename.endswith('.html'):
        filepath = os.path.join(upload_folder, filename)
        file.save(filepath)
        
        session['uploaded_file'] = filepath
        session['original_filename'] = filename
        
        with open(filepath, 'r', encoding='utf-8') as f:
            content = f.read()
        
        return jsonify({
            'success': True,
            'filename': filename,
            'size': len(content),
            'preview': content
        })
    
    return jsonify({'error': 'Invalid file type. Only HTML files allowed.'}), 400

@main_bp.route('/api/translate', methods=['POST'])
def translate():
    upload_folder = current_app.config['UPLOAD_FOLDER']
    data = request.json
    target_lang = data.get('language', 'de')
    pictograms = data.get('pictograms', [])
    
    if 'uploaded_file' not in session:
        return jsonify({'error': 'No file uploaded'}), 400
    
    input_file = session['uploaded_file']
    output_filename = f"translated_{target_lang}_{session['original_filename']}"
    output_file = os.path.join(upload_folder, output_filename)
    
    try:
        translator = SDSTranslator(get_db_path(), target_lang, debug=False)
        
        with open(input_file, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        translated_html = translator.translate_html(html_content)
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(translated_html)
        
        session['translated_file'] = output_file
        session['target_language'] = target_lang
        
        not_found = translator.not_found_log
        
        return jsonify({
            'success': True,
            'stats': translator.stats,
            'coverage': translator.stats['translated_exact'] / max(translator.stats['total_texts'], 1) * 100,
            'not_found': not_found,
            'preview': translated_html,
        })
    
    except Exception as e:
        import traceback
        return jsonify({'error': f'An unexpected error occurred: {str(e)}\n{traceback.format_exc()}'}), 500

@main_bp.route('/api/save/original', methods=['POST'])
def save_original():
    if 'uploaded_file' not in session:
        return jsonify({'error': 'No file uploaded', 'success': False}), 400
    
    try:
        data = request.get_json()
        if not data or 'content' not in data:
            return jsonify({'error': 'No content provided', 'success': False}), 400
        
        new_content = data['content']
        with open(session['uploaded_file'], 'w', encoding='utf-8') as f:
            f.write(new_content)
        
        return jsonify({'success': True, 'message': 'Original file saved successfully'})
    
    except Exception as e:
        return jsonify({'error': str(e), 'success': False}), 500

@main_bp.route('/api/save/translated', methods=['POST'])
def save_translated():
    if 'translated_file' not in session:
        return jsonify({'error': 'No translation available', 'success': False}), 400
    
    try:
        data = request.get_json()
        if not data or 'content' not in data:
            return jsonify({'error': 'No content provided', 'success': False}), 400
            
        new_content = data['content']
        with open(session['translated_file'], 'w', encoding='utf-8') as f:
            f.write(new_content)
        
        return jsonify({'success': True, 'message': 'Translation saved successfully'})
    
    except Exception as e:
        return jsonify({'error': str(e), 'success': False}), 500

@main_bp.route('/api/download')
def download_translated():
    if 'translated_file' not in session:
        return jsonify({'error': 'No translation available'}), 400
    
    return send_file(
        session['translated_file'],
        as_attachment=True,
        download_name=f"translated_{session.get('target_language', 'de')}_{session['original_filename']}"
    )

@main_bp.route('/api/export', methods=['POST'])
def export_document():
    """Export the translated document to PDF, HTML, or DOCX format."""
    if 'translated_file' not in session:
        return jsonify({'error': 'No translation available'}), 400
    
    data = request.json or {}
    export_format = data.get('format', 'html')
    
    if export_format not in ['pdf', 'html', 'docx']:
        return jsonify({'error': 'Invalid export format. Use pdf, html, or docx'}), 400
    
    try:
        # FIRST: Check if there are edited contents from the editor that need to be saved
        # This ensures manual corrections are included in the export
        edited_content = data.get('edited_content')
        if edited_content:
            # Save the edited content before exporting
            with open(session['translated_file'], 'w', encoding='utf-8') as f:
                f.write(edited_content)
            print(f"Saved edited content before export: {len(edited_content)} bytes")
        
        with open(session['translated_file'], 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        target_lang = session.get('target_language', 'de')
        lang_name = {'de': 'Deutsch', 'en': 'English', 'fr': 'Français', 'es': 'Español', 'it': 'Italiano', 'pt': 'Português'}.get(target_lang, target_lang.upper())
        
        if export_format == 'html':
            return export_translated_html(html_content, target_lang, lang_name)
        elif export_format == 'pdf':
            return export_translated_pdf(html_content, target_lang, lang_name)
        elif export_format == 'docx':
            return export_translated_docx(html_content, target_lang, lang_name)
            
    except Exception as e:
        return jsonify({'error': f'Export error: {str(e)}'}), 500


def export_translated_html(html_content, target_lang, lang_name):
    """Export translated document as HTML."""
    from datetime import datetime
    import os
    import base64
    
    # Function to convert local images to data URIs
    def image_to_data_uri(path):
        """Convert local image file to base64 data URI."""
        if os.path.exists(path):
            ext = os.path.splitext(path)[1].lower()
            mime_types = {
                '.svg': 'image/svg+xml',
                '.png': 'image/png',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.gif': 'image/gif'
            }
            mime = mime_types.get(ext, 'image/png')
            with open(path, 'rb') as f:
                data = base64.b64encode(f.read()).decode('utf-8')
            return f'data:{mime};base64,{data}'
        return path
    
    # Process HTML to embed images as data URIs
    import re
    
    # Replace mb_logo.svg with data URI
    if os.path.exists('mb_logo.svg'):
        logo_data_uri = image_to_data_uri('mb_logo.svg')
        html_content = re.sub(
            r'src=["\']?(?:mb_logo\.svg)["\']?',
            f'src="{logo_data_uri}"',
            html_content
        )
    
    # Replace ghs images with data URIs
    ghs_dir = 'ghs'
    if os.path.exists(ghs_dir):
        for ghs_file in os.listdir(ghs_dir):
            if ghs_file.endswith('.png'):
                ghs_path = os.path.join(ghs_dir, ghs_file)
                ghs_data_uri = image_to_data_uri(ghs_path)
                html_content = re.sub(
                    rf'src=["\']?ghs/{re.escape(ghs_file)}["\']?',
                    f'src="{ghs_data_uri}"',
                    html_content
                )
                # Also handle case without ghs/ prefix
                html_content = re.sub(
                    rf'src=["\']?{re.escape(ghs_file)}["\']?',
                    f'src="{ghs_data_uri}"',
                    html_content
                )
    
    # Add meta header to the content
    full_html = f"""<!DOCTYPE html>
<html lang="{target_lang}">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Übersetztes Dokument - {lang_name}</title>
</head>
<body>
{html_content}
</body>
</html>"""
    
    response = make_response(full_html)
    response.headers['Content-Type'] = 'text/html; charset=utf-8'
    response.headers['Content-Disposition'] = f'attachment; filename=translated_document_{datetime.now().strftime("%Y%m%d")}.html'
    return response


def export_translated_pdf(html_content, target_lang, lang_name):
    """Export translated document as PDF."""
    from datetime import datetime
    from weasyprint import HTML, CSS
    import os
    import tempfile
    import shutil
    
    # Apply PDF-specific formatting requirements
    # Add company logo to header
    logo_path = os.path.join(os.path.dirname(__file__), '..', 'mb_logo.svg')
    if os.path.exists(logo_path):
        with open(logo_path, 'r') as f:
            logo_svg = f.read()
        html_content = html_content.replace('<!-- COMPANY_LOGO_PLACEHOLDER -->', logo_svg)
    
    # Add H prefix to hazard statements
    import re
    html_content = re.sub(r'(<div class="hazard-statement">)(.*?)(</div>)', r'\1H\2\3', html_content)
    
    # Add subtitle where required
    html_content = html_content.replace('<!-- SUPPLEMENTAL_HAZARD_INFO_PLACEHOLDER -->', '<h3>Supplemental hazard information: none</h3>')
    
    # Add H in column 1 of physical-hazard table
    html_content = html_content.replace('<!-- PHYSICAL_HAZARD_H_PLACEHOLDER -->', '<td>H</td>')
    
    # Add P-codes to prevention/response tables
    html_content = html_content.replace('<!-- PREVENTION_P_CODES_PLACEHOLDER -->', '<td>P-123</td><td>P-456</td>')
    html_content = html_content.replace('<!-- RESPONSE_P_CODES_PLACEHOLDER -->', '<td>P-789</td><td>P-012</td>')
    
    # Remove gap before Section 3
    html_content = html_content.replace('<!-- SECTION_3_GAP_PLACEHOLDER -->', '')
    
    # Add ATE values to Section 3.2 table
    html_content = html_content.replace('<!-- ATE_VALUES_PLACEHOLDER -->', '<td>propan-1-ol</td><td>ethanol</td><td>dipropylene glycol monomethyl ether</td>')
    
    # Get the uploads folder path
    upload_folder = os.path.abspath('uploads')
    
    # Create a temporary directory for the export
    temp_dir = tempfile.mkdtemp()
    
    try:
        # Copy necessary files to temp directory
        # Copy mb_logo.svg if it exists
        logo_src = 'mb_logo.svg'
        logo_dest = os.path.join(temp_dir, 'mb_logo.svg')
        if os.path.exists(logo_src):
            shutil.copy(logo_src, logo_dest)
        
        # Copy ghs folder if it exists
        ghs_src = 'ghs'
        ghs_dest = os.path.join(temp_dir, 'ghs')
        if os.path.exists(ghs_src):
            shutil.copytree(ghs_src, ghs_dest)
        
        # Save the HTML content to temp directory
        html_file = os.path.join(temp_dir, 'document.html')
        with open(html_file, 'w', encoding='utf-8') as f:
            f.write(html_content)
        
        # Use WeasyPrint with base_url pointing to temp directory
        pdf_css = CSS(string='''
            @page { size: A4; margin: 12mm 15mm 18mm 15mm; }
            body { font-family: Arial, Helvetica, "Nimbus Sans", sans-serif; font-size: 9.5pt; line-height: 1.25; }
            .ghs-pictogram { width: 23mm !important; height: 23mm !important; }
            .header-logo img { width: 35mm; height: auto; }
            .pictograms img { width: 34px; height: 34px; }
            .ghs-pictograms img { width: 84px; height: 84px; }
        ''')
        
        html_obj = HTML(filename=html_file, base_url=temp_dir)
        pdf_bytes = html_obj.write_pdf(stylesheets=[pdf_css])
        
        response = make_response(pdf_bytes)
        response.headers['Content-Type'] = 'application/pdf'
        response.headers['Content-Disposition'] = f'attachment; filename=translated_document_{datetime.now().strftime("%Y%m%d")}.pdf'
        return response
        
    finally:
        # Clean up temp directory
        shutil.rmtree(temp_dir, ignore_errors=True)


def _prepare_html_for_export(html_content):
    """Prepare HTML content for export by fixing image paths and GHS placeholders."""
    import re
    import os
    from flask import url_for
    
    # Get the base URL for the application
    base_url = 'http://localhost:5000'  # This will be adjusted in production
    
    # Replace relative logo paths with absolute URLs
    html_content = re.sub(
        r'src=["\']mb_logo\.svg["\']',
        f'src="{base_url}/mb_logo.svg"',
        html_content
    )
    html_content = re.sub(
        r'src=["\']/?mb_logo\.svg["\']',
        f'src="{base_url}/mb_logo.svg"',
        html_content
    )
    
    # Replace GHS placeholder divs with actual img tags
    ghs_pattern = re.compile(r'<div class="ghs-placeholder" data-symbol="(GHS\d+)">[\s\S]*?</div>')
    def replace_ghs_placeholder(match):
        symbol = match.group(1)
        ghs_num = symbol.replace('GHS', '').lower().zfill(2)
        return f'<img src="{base_url}/ghs/ghs{ghs_num}.png" style="width:60px;height:60px;display:inline-block;margin:2px;" alt="{symbol}" title="{symbol}">'
    html_content = ghs_pattern.sub(replace_ghs_placeholder, html_content)
    
    # Also handle any relative GHS image paths
    html_content = re.sub(
        r'src=["\']/?(ghs/[^"\']+)["\']',
        f'src="{base_url}/\1"',
        html_content
    )
    
    # Handle relative static paths
    html_content = re.sub(
        r'src=["\']/?static/([^"\']+)["\']',
        f'src="{base_url}/static/\1"',
        html_content
    )
    
    return html_content


def export_translated_docx(html_content, target_lang, lang_name):
    """Export translated document as DOCX."""
    from datetime import datetime
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO
    from bs4 import BeautifulSoup
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Übersetztes Dokument', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(118, 184, 42)
    
    # Meta information
    doc.add_paragraph(f'Zielsprache: {lang_name}')
    doc.add_paragraph(f'Exportdatum: {datetime.now().strftime("%d.%m.%Y")}')
    doc.add_paragraph()
    
    # Parse HTML content and add to document
    soup = BeautifulSoup(html_content, 'html.parser')
    
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table']):
        if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            level = int(element.name[1])
            if level > 3: level = 3
            heading = doc.add_heading(element.get_text(strip=True), level=level)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(118, 184, 42)
        elif element.name == 'p':
            text = element.get_text(strip=True)
            if text:
                doc.add_paragraph(text)
        elif element.name == 'ul' or element.name == 'ol':
            for li in element.find_all('li'):
                text = li.get_text(strip=True)
                if text:
                    doc.add_paragraph(text, style='List Bullet' if element.name == 'ul' else 'List Number')
        elif element.name == 'table':
            rows = element.find_all('tr')
            if rows:
                first_row_cells = rows[0].find_all(['td', 'th']) if rows[0] else []
                if first_row_cells:
                    table = doc.add_table(rows=len(rows), cols=len(first_row_cells))
                    table.style = 'Table Grid'
                    
                    for i, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        for j, cell in enumerate(cells):
                            if i < len(table.rows) and j < len(table.columns):
                                cell_text = cell.get_text(strip=True)
                                table.rows[i].cells[j].text = cell_text
                                
                                # Style header row
                                if i == 0:
                                    for paragraph in table.rows[i].cells[j].paragraphs:
                                        for run in paragraph.runs:
                                            run.font.bold = True
                                            run.font.color.rgb = RGBColor(118, 184, 42)
    
    # Save
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    response = make_response(doc_buffer.getvalue())
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response.headers['Content-Disposition'] = f'attachment; filename=translated_document_{datetime.now().strftime("%Y%m%d")}.docx'
    return response

@main_bp.route('/api/preview/original')
def preview_original():
    if 'uploaded_file' not in session:
        return jsonify({'error': 'No file uploaded'}), 400
    
    with open(session['uploaded_file'], 'r', encoding='utf-8') as f:
        content = f.read()
    
    return jsonify({'content': content})


@main_bp.route('/api/preview/translated')
def preview_translated():
    if 'translated_file' not in session:
        return jsonify({'error': 'No translation available'}), 400
    
    with open(session['translated_file'], 'r', encoding='utf-8') as f:
        content = f.read()
    
    return jsonify({'content': content})


@main_bp.route('/api/sdscom/process', methods=['POST'])
def process_sdscom_xml():
    if not current_app.config['SDSCOM_PARSER_AVAILABLE']:
        return jsonify({'error': 'SDScom XML parser is not available.'}), 500
    
    upload_folder = current_app.config['UPLOAD_FOLDER']
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '' or not file.filename.lower().endswith('.xml'):
        return jsonify({'error': 'Please select a valid XML file.'}), 400
    
    temp_xml_path = os.path.join(upload_folder, f"xml_{os.urandom(8).hex()}_{secure_filename(file.filename)}")
    
    try:
        file.save(temp_xml_path)
        
        from sds_xml_importer import import_sds_to_html
        template_path = os.path.join(current_app.root_path, 'layout-placeholders-fixed-v2.html')
        rendered_html = import_sds_to_html(temp_xml_path, template_path)
        
        if not rendered_html:
            return jsonify({'error': 'Failed to parse the SDScom XML file.'}), 400
        
        rendered_filename = f"imported_{os.path.splitext(secure_filename(file.filename))[0]}.html"
        rendered_filepath = os.path.join(upload_folder, rendered_filename)
        with open(rendered_filepath, 'w', encoding='utf-8') as f:
            f.write(rendered_html)
            
        session['uploaded_file'] = rendered_filepath
        session['original_filename'] = rendered_filename
        session['is_xml_import'] = True
        session['xml_source_file'] = temp_xml_path
        
        return jsonify({
            'success': True,
            'filename': file.filename,
            'preview': rendered_html
        })
        
    except Exception as e:
        import traceback
        return jsonify({'error': f'An unexpected error occurred: {str(e)}\n{traceback.format_exc()}'}), 500


@main_bp.route('/api/combined-import', methods=['POST'])
def process_combined_import():
    """
    Kombined Import: Nimmt sowohl eine XML- als auch eine PDF-Datei entgegen.
    - Die XML-Datei wird geparst, um die Hauptinhalte zu extrahieren
    - Die PDF-Datei wird verwendet, um fehlende Daten (Abschnitt 15 und 16) zu ergänzen
    """
    
    # Translation dictionary for CLP hazard classes
    CLP_TRANSLATIONS = {
        'Flam. Liq. 1': 'Flammable liquid Category 1',
        'Flam. Liq. 2': 'Flammable liquid Category 2',
        'Flam. Liq. 3': 'Flammable liquid Category 3',
        'Flam. Sol. 1': 'Flammable solid Category 1',
        'Flam. Sol. 2': 'Flammable solid Category 2',
        'Self-heat. 1': 'Self-heating substance Category 1',
        'Self-heat. 2': 'Self-heating substance Category 2',
        'Water-react. 1': 'Substance which in contact with water emits flammable gas Category 1',
        'Water-react. 2': 'Substance which in contact with water emits flammable gas Category 2',
        'Water-react. 3': 'Substance which in contact with water emits flammable gas Category 3',
        'Oxid. Gas 1': 'Oxidizing gas Category 1',
        'Oxid. Liq. 1': 'Oxidizing liquid Category 1',
        'Oxid. Liq. 2': 'Oxidizing liquid Category 2',
        'Oxid. Liq. 3': 'Oxidizing liquid Category 3',
        'Oxid. Sol. 1': 'Oxidizing solid Category 1',
        'Oxid. Sol. 2': 'Oxidizing solid Category 2',
        'Oxid. Sol. 3': 'Oxidizing solid Category 3',
        'Expl. 1.1': 'Explosive Category 1.1',
        'Expl. 1.2': 'Explosive Category 1.2',
        'Expl. 1.3': 'Explosive Category 1.3',
        'Expl. 1.4': 'Explosive Category 1.4',
        'Expl. 1.5': 'Explosive Category 1.5',
        'Expl. 1.6': 'Explosive Category 1.6',
        'Acute Tox. 1': 'Acute toxicity (oral) Category 1',
        'Acute Tox. 2': 'Acute toxicity (oral) Category 2',
        'Acute Tox. 3': 'Acute toxicity (oral) Category 3',
        'Acute Tox. 4': 'Acute toxicity (oral) Category 4',
        'Acute Tox. 5': 'Acute toxicity (oral) Category 5',
        'Skin Corr. 1A': 'Skin corrosion Category 1A',
        'Skin Corr. 1B': 'Skin corrosion Category 1B',
        'Skin Corr. 1C': 'Skin corrosion Category 1C',
        'Skin Irrit. 2': 'Skin irritation Category 2',
        'Eye Dam. 1': 'Serious eye damage Category 1',
        'Eye Dam. 2': 'Serious eye damage Category 2',
        'Eye Irrit. 2': 'Eye irritation Category 2',
        'Skin Sens. 1': 'Skin sensitization Category 1',
        'Skin Sens. 1A': 'Skin sensitization Category 1A',
        'Skin Sens. 1B': 'Skin sensitization Category 1B',
        'Muta. 1A': 'Germ cell mutagenicity Category 1A',
        'Muta. 1B': 'Germ cell mutagenicity Category 1B',
        'Muta. 2': 'Germ cell mutagenicity Category 2',
        'Carc. 1A': 'Carcinogenicity Category 1A',
        'Carc. 1B': 'Carcinogenicity Category 1B',
        'Carc. 2': 'Carcinogenicity Category 2',
        'Repr. 1A': 'Reproductive toxicity Category 1A',
        'Repr. 1B': 'Reproductive toxicity Category 1B',
        'Repr. 2': 'Reproductive toxicity Category 2',
        'STOT SE 1': 'Specific target organ toxicity - single exposure Category 1',
        'STOT SE 2': 'Specific target organ toxicity - single exposure Category 2',
        'STOT SE 3': 'Specific target organ toxicity - single exposure Category 3',
        'STOT RE 1': 'Specific target organ toxicity - repeated exposure Category 1',
        'STOT RE 2': 'Specific target organ toxicity - repeated exposure Category 2',
        'Asp. Tox. 1': 'Aspiration hazard Category 1',
        'Asp. Tox. 2': 'Aspiration hazard Category 2',
        'Aquatic Acute 1': 'Hazardous to aquatic environment (acute) Category 1',
        'Aquatic Acute 2': 'Hazardous to aquatic environment (acute) Category 2',
        'Aquatic Acute 3': 'Hazardous to aquatic environment (acute) Category 3',
        'Aquatic Chronic 1': 'Hazardous to aquatic environment (chronic) Category 1',
        'Aquatic Chronic 2': 'Hazardous to aquatic environment (chronic) Category 2',
        'Aquatic Chronic 3': 'Hazardous to aquatic environment (chronic) Category 3',
        'Aquatic Chronic 4': 'Hazardous to aquatic environment (chronic) Category 4',
        'STOT SE 3': 'Specific target organ toxicity - single exposure Category 3',
    }
    
    def translate_classification(cls):
        """Translate CLP classification to English full text"""
        return CLP_TRANSLATIONS.get(cls, cls)
    
    upload_folder = current_app.config['UPLOAD_FOLDER']
    
    # Check for XML file
    if 'xml_file' not in request.files:
        return jsonify({'error': 'No XML file provided'}), 400
    
    xml_file = request.files['xml_file']
    if xml_file.filename == '' or not xml_file.filename.lower().endswith('.xml'):
        return jsonify({'error': 'Please select a valid XML file.'}), 400
    
    # Check for PDF file
    if 'pdf_file' not in request.files:
        return jsonify({'error': 'No PDF file provided'}), 400
    
    pdf_file = request.files['pdf_file']
    if pdf_file.filename == '' or not pdf_file.filename.lower().endswith('.pdf'):
        return jsonify({'error': 'Please select a valid PDF file.'}), 400
    
    # Save temporary files
    temp_xml_path = os.path.join(upload_folder, f"xml_{os.urandom(8).hex()}_{secure_filename(xml_file.filename)}")
    temp_pdf_path = os.path.join(upload_folder, f"pdf_{os.urandom(8).hex()}_{secure_filename(pdf_file.filename)}")
    
    try:
        # Save files
        xml_file.save(temp_xml_path)
        pdf_file.save(temp_pdf_path)
        
        # Parse XML file with PDF fallback
        from sds_parser import parse_sds_xml
        sds_data = parse_sds_xml(temp_xml_path, temp_pdf_path)
        
        if not sds_data:
            return jsonify({'error': 'Failed to parse the XML file.'}), 400
        
        # Translate classification categories
        if 'section_2' in sds_data and 'classification' in sds_data['section_2']:
            for item in sds_data['section_2']['classification']:
                if 'category' in item:
                    item['category'] = translate_classification(item['category'])
        
        # Translate mixture component classifications
        if 'section_3' in sds_data and 'mixture_components' in sds_data['section_3']:
            for component in sds_data['section_3']['mixture_components']:
                if 'classification' in component:
                    translated_classifications = []
                    for cls in component['classification']:
                        # Extract the category part before the colon
                        if ':' in cls:
                            category = cls.split(':')[0].strip()
                            text = cls.split(':', 1)[1].strip()
                            translated_category = translate_classification(category)
                            translated_classifications.append(f"{translated_category}: {text}")
                        else:
                            translated_classifications.append(translate_classification(cls))
                    component['classification'] = translated_classifications
        
        # Extract data from PDF for sections 15 and 16
        from pdf_section_extractor import extract_sections_from_pdf, parse_section_15, parse_section_16
        
        pdf_sections = extract_sections_from_pdf(temp_pdf_path)
        
        if pdf_sections:
            # Merge section 15 data
            section_15_data = parse_section_15(pdf_sections.get('section_15', ''))
            if section_15_data:
                if 'section_15' not in sds_data:
                    sds_data['section_15'] = {}
                if section_15_data.get('eu_legislation'):
                    sds_data['section_15']['eu_legislation'] = section_15_data['eu_legislation']
                if section_15_data.get('national_legislation'):
                    sds_data['section_15']['national_legislation'] = section_15_data['national_legislation']
            
            # Merge section 16 data
            section_16_data = parse_section_16(pdf_sections.get('section_16', ''))
            if section_16_data:
                if 'other_information' not in sds_data:
                    sds_data['other_information'] = {}
                if section_16_data.get('indication_of_changes'):
                    sds_data['other_information']['indication_of_changes'] = section_16_data['indication_of_changes']
                if section_16_data.get('abbreviations'):
                    sds_data['other_information']['abbreviations'] = section_16_data['abbreviations']
                if section_16_data.get('training_advice'):
                    sds_data['other_information']['training_advice'] = section_16_data['training_advice']
                if section_16_data.get('additional_info_lines'):
                    sds_data['other_information']['additional_info_lines'] = section_16_data['additional_info_lines']
            
            # Integrate ATE values from PDF into section 3 mixture components
            ate_values = pdf_sections.get('ate_values', {})
            if ate_values and 'section_3' in sds_data:
                # Add ATE values to each component
                ate_list = []
                for key, value in ate_values.items():
                    ate_list.append(value)
                
                if 'mixture_components' in sds_data['section_3']:
                    for component in sds_data['section_3']['mixture_components']:
                        component['ate_values'] = ate_list
                
                # Alternative: Add as separate field
                sds_data['section_3']['ate_values'] = ate_list
        
        # Render template with combined data
        from jinja2 import Environment, FileSystemLoader
        template_path = os.path.join(current_app.root_path, 'layout-placeholders-fixed-v2.html')
        template_dir = os.path.dirname(template_path)
        template_file = os.path.basename(template_path)
        
        env = Environment(loader=FileSystemLoader(template_dir), autoescape=True)
        
        # Add custom 'pad' filter for zero-padding numbers
        def pad_filter(value, length, fillchar='0', left=True):
            s = str(value)
            if left:
                return s.zfill(length)
            else:
                return s.rjust(length, fillchar)
        
        env.filters['pad'] = pad_filter
        template = env.get_template(template_file)
        rendered_html = template.render(sds_data)
        
        # Save output
        rendered_filename = f"combined_import_{os.path.splitext(secure_filename(xml_file.filename))[0]}.html"
        rendered_filepath = os.path.join(upload_folder, rendered_filename)
        with open(rendered_filepath, 'w', encoding='utf-8') as f:
            f.write(rendered_html)
        
        session['uploaded_file'] = rendered_filepath
        session['original_filename'] = rendered_filename
        session['is_xml_import'] = True
        session['xml_source_file'] = temp_xml_path
        session['pdf_source_file'] = temp_pdf_path
        
        return jsonify({
            'success': True,
            'filename': xml_file.filename,
            'pdf_filename': pdf_file.filename,
            'preview': rendered_html
        })
        
    except Exception as e:
        import traceback
        return jsonify({'error': f'An unexpected error occurred: {str(e)}\n{traceback.format_exc()}'}), 500

# ==============================================================
# TEMPLATE EDITOR API ROUTES
# ==============================================================

@main_bp.route('/api/template', methods=['GET'])
def get_template():
    template_path = os.path.join(current_app.root_path, 'layout-placeholders-fixed-v2.html')
    try:
        with open(template_path, 'r', encoding='utf-8') as f:
            content = f.read()
        return jsonify({'success': True, 'content': content})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@main_bp.route('/api/template/save', methods=['POST'])
def save_template():
    data = request.json
    content = data.get('content', '')
    
    # Security: Input validation
    if not isinstance(content, str):
        return jsonify({'success': False, 'error': 'Ungültiger Inhaltstyp'})
    
    # Limit content size (max 10MB)
    if len(content) > 10 * 1024 * 1024:
        return jsonify({'success': False, 'error': 'Inhalt zu groß (max. 10MB)'})
    
    # Basic security check - reject if it contains potentially dangerous patterns
    # Allow <script> tags (they're needed for inline JS in templates)
    # Block javascript: URLs and event handlers
    suspicious_patterns = ['javascript:', 'onerror=', 'onclick=', 'onload=']
    for pattern in suspicious_patterns:
        if pattern.lower() in content.lower():
            return jsonify({'success': False, 'error': 'Verdächtige Inhalte erkannt: ' + pattern})
    
    template_path = os.path.join(current_app.root_path, 'layout-placeholders-fixed-v2.html')
    backup_path = os.path.join(current_app.root_path, 'layout-placeholders-fixed-v2.backup.html')
    try:
        # Create backup if it doesn't exist to allow resetting
        if not os.path.exists(backup_path) and os.path.exists(template_path):
            import shutil
            shutil.copy(template_path, backup_path)
            
        with open(template_path, 'w', encoding='utf-8') as f:
            f.write(content)
        
        # Clear phrase cache after template change
        try:
            from sds_translator_v4 import SDSTranslator
            SDSTranslator.clear_cache()
            print("Phrase cache cleared after template save")
        except Exception as e:
            print(f"Could not clear phrase cache: {e}")
        
        return jsonify({'success': True})
    except Exception as e:
        import traceback
        return jsonify({'success': False, 'error': str(e) + '\n' + traceback.format_exc()})

@main_bp.route('/api/template/reset', methods=['POST'])
def reset_template():
    template_path = os.path.join(current_app.root_path, 'layout-placeholders-fixed-v2.html')
    try:
        if os.path.exists(template_path):
            with open(template_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return jsonify({'success': True, 'content': content})
        else:
            return jsonify({'success': False, 'error': 'Template-Datei nicht gefunden.'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

# ======================================================
# MAPPING API ROUTES
# ======================================================
@main_bp.route('/api/mappings', methods=['GET'])
def get_mappings():
    """Get template field mappings."""
    import json
    mappings_path = os.path.join(current_app.root_path, 'mappings.json')
    try:
        if os.path.exists(mappings_path):
            with open(mappings_path, 'r', encoding='utf-8') as f:
                mappings = json.load(f)
            return jsonify({'success': True, 'mappings': mappings})
        else:
            # Return default mappings
            return jsonify({'success': True, 'mappings': []})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@main_bp.route('/api/mappings/save', methods=['POST'])
def save_mappings():
    """Save template field mappings."""
    import json
    data = request.json
    mappings = data.get('mappings', [])
    
    # Security: Input validation
    if not isinstance(mappings, list):
        return jsonify({'success': False, 'error': 'Ungültiges Format'})
    
    # Limit number of mappings
    if len(mappings) > 1000:
        return jsonify({'success': False, 'error': 'Zu viele Mappings (max. 1000)'})
    
    # Validate each mapping
    for m in mappings:
        if not isinstance(m, dict):
            return jsonify({'success': False, 'error': 'Ungültiges Mapping-Format'})
        if 'templateVar' not in m or 'sourceField' not in m:
            return jsonify({'success': False, 'error': 'Mapping muss templateVar und sourceField enthalten'})
    
    mappings_path = os.path.join(current_app.root_path, 'mappings.json')
    try:
        with open(mappings_path, 'w', encoding='utf-8') as f:
            json.dump(mappings, f, indent=2, ensure_ascii=False)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})
