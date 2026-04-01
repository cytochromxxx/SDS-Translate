from flask import Blueprint, jsonify, request, session, Response, make_response
import sqlite3
import uuid
from database import get_db_connection, get_current_db_info, get_available_databases, set_db_path, ensure_database_indices, get_database_stats
from utils import AVAILABLE_LANGUAGES, LANG_TO_COLUMN, parse_flag_format

database_bp = Blueprint('database', __name__)

@database_bp.route('/api/databases', methods=['GET'])
def get_databases_route():
    return jsonify({
        'current': get_current_db_info(),
        'available': get_available_databases()
    })

@database_bp.route('/api/databases/select', methods=['POST'])
def select_database_route():
    data = request.json
    db_key = data.get('database')
    
    if not db_key:
        return jsonify({'error': 'No database key provided'}), 400
    
    success, message = set_db_path(db_key)
    
    if success:
        session['selected_database'] = db_key
        # Clear cache when switching databases
        try:
            from sds_translator_v4 import SDSTranslator
            SDSTranslator.clear_cache()
            print("Phrase cache cleared after database switch")
        except Exception as e:
            print(f"Could not clear phrase cache: {e}")
        return jsonify({
            'success': True,
            'message': message,
            'current': get_current_db_info()
        })
    else:
        return jsonify({
            'success': False,
            'error': message,
            'current': get_current_db_info()
        }), 400

@database_bp.route('/api/databases/indices', methods=['POST'])
def create_indices_route():
    """
    Create database indices for better search performance.
    This is safe to call multiple times - indices are only created if they don't exist.
    """
    result = ensure_database_indices()
    if result.get('success'):
        return jsonify({
            'success': True,
            'message': f"Created {result.get('indices_created', 0)} indices",
            'indices': result.get('indices', []),
            'errors': result.get('errors', [])
        })
    else:
        return jsonify({'error': result.get('error', 'Failed to create indices')}), 500

@database_bp.route('/api/databases/stats', methods=['GET'])
def get_db_stats_route():
    """Get database statistics including index information."""
    result = get_database_stats()
    if result.get('success'):
        return jsonify(result)
    else:
        return jsonify({'error': result.get('error', 'Failed to get stats')}), 500

@database_bp.route('/api/phrases/search', methods=['POST'])
def search_phrases_route():
    data = request.json or {}
    query = data.get('q', '')
    lang = data.get('lang', 'de')
    page = data.get('page', 1)
    per_page = 50
    sort_by = data.get('sort_by', 'id_desc')  # id_asc, id_desc, en_asc, en_desc
    filter_en = data.get('filter_en', '')  # Filter by English text

    # Validate language code
    if lang not in AVAILABLE_LANGUAGES:
        return jsonify({'error': 'Invalid language code'}), 400

    # Validate and sanitize lang_col - only allow known columns
    lang_col = f"{lang}_original"
    
    # Whitelist validation for sort_by to prevent SQL injection
    allowed_sort_options = {
        'id_asc': 'ORDER BY id ASC',
        'id_desc': 'ORDER BY id DESC',
        'en_asc': 'ORDER BY en_original ASC',
        'en_desc': 'ORDER BY en_original DESC'
    }
    order_by = allowed_sort_options.get(sort_by, 'ORDER BY id DESC')

    # Build WHERE clause for English filter
    where_clauses = []
    params = []
    
    if query:
        search_pattern = f'%{query}%'
        where_clauses.append("(en_original LIKE ? OR en_original LIKE ?)")  # Fixed: use en_original twice instead of dynamic column
        params.extend([search_pattern, search_pattern])
    
    if filter_en:
        filter_pattern = f'%{filter_en}%'
        where_clauses.append("en_original LIKE ?")
        params.extend([filter_pattern])
    
    where_sql = " AND ".join(where_clauses) if where_clauses else "1=1"

    conn = get_db_connection()
    
    try:
        count_cursor = conn.execute(
            f"SELECT COUNT(*) FROM phrases WHERE {where_sql}",
            params
        )
        total_phrases = count_cursor.fetchone()[0]

        cursor = conn.execute(
            f"""SELECT id, en_original, {lang_col} as translation
                FROM phrases
                WHERE {where_sql}
                {order_by}
                LIMIT ? OFFSET ?""",
            params + [per_page, (page - 1) * per_page]
        )
        
        phrases = [dict(row) for row in cursor.fetchall()]
        conn.close()
        
        return jsonify({
            'phrases': phrases,
            'total': total_phrases,
            'page': page,
            'per_page': per_page
        })
    
    except Exception as e:
        conn.close()
        return jsonify({'error': str(e)}), 500

@database_bp.route('/api/phrases/<phrase_id>/full', methods=['GET'])
def get_phrase_full_route(phrase_id):
    conn = get_db_connection()
    
    cursor = conn.execute(
        """SELECT * FROM phrases WHERE id = ?""",
        (phrase_id,)
    )
    phrase = cursor.fetchone()
    conn.close()
    
    if phrase:
        return jsonify(dict(phrase))
    return jsonify({'error': 'Phrase not found'}), 404

@database_bp.route('/api/phrases/<phrase_id>', methods=['PUT', 'DELETE'])
def manage_phrase_route(phrase_id):
    conn = get_db_connection()
    
    if request.method == 'PUT':
        data = request.json
        update_fields = []
        values = []
        
        # Explicitly handle English original text (since 'en' is not in AVAILABLE_LANGUAGES)
        if 'en_original' in data:
            update_fields.append("en_original = ?")
            values.append(data['en_original'])
            
        for lang_code in AVAILABLE_LANGUAGES:
            field = f'{lang_code}_original'
            if field in data:
                update_fields.append(f"{field} = ?")
                values.append(data[field])
        
        if not update_fields:
            return jsonify({'error': 'No fields to update'}), 400
        
        values.append(phrase_id)
        query = f"UPDATE phrases SET {', '.join(update_fields)} WHERE id = ?"
        
        try:
            conn.execute(query, values)
            conn.commit()
        except Exception as e:
            return jsonify({'error': f'Database error: {str(e)}'}), 500
        finally:
            conn.close()
        # Clear phrase cache after update
        try:
            from sds_translator_v4 import SDSTranslator
            SDSTranslator.clear_cache()
        except Exception as e:
            print(f"Could not clear phrase cache: {e}")
        return jsonify({'success': True, 'message': 'Phrase updated'})
    
    elif request.method == 'DELETE':
        try:
            conn.execute("DELETE FROM phrases WHERE id = ?", (phrase_id,))
            conn.commit()
        except Exception as e:
            return jsonify({'error': f'Database error: {str(e)}'}), 500
        finally:
            conn.close()
        # Clear phrase cache after delete
        try:
            from sds_translator_v4 import SDSTranslator
            SDSTranslator.clear_cache()
        except Exception as e:
            print(f"Could not clear phrase cache: {e}")
        return jsonify({'success': True, 'message': 'Phrase deleted'})

@database_bp.route('/api/phrases/bulk/delete', methods=['POST'])
def bulk_delete_phrases_route():
    """Delete multiple phrases at once."""
    data = request.json
    phrase_ids = data.get('ids', [])
    
    if not phrase_ids:
        return jsonify({'error': 'No phrase IDs provided'}), 400
    
    if not isinstance(phrase_ids, list):
        return jsonify({'error': 'ids must be an array'}), 400
    
    conn = get_db_connection()
    
    try:
        # Create placeholders for the query
        placeholders = ', '.join(['?'] * len(phrase_ids))
        cursor = conn.execute(f"DELETE FROM phrases WHERE id IN ({placeholders})", phrase_ids)
        conn.commit()
        deleted_count = cursor.rowcount
        conn.close()
        
        # Clear phrase cache after bulk delete
        try:
            from sds_translator_v4 import SDSTranslator
            SDSTranslator.clear_cache()
        except Exception as e:
            print(f"Could not clear phrase cache: {e}")
        
        return jsonify({
            'success': True,
            'message': f'{deleted_count} phrase(s) deleted',
            'deleted_count': deleted_count
        })
    
    except Exception as e:
        conn.close()
        return jsonify({'error': f'Database error: {str(e)}'}), 500

@database_bp.route('/api/phrases', methods=['POST'])
def add_phrase_route():
    data = request.json
    
    en_text = data.get('en_original', '').strip()
    if not en_text:
        return jsonify({'error': 'English text is required'}), 400
    
    conn = get_db_connection()
    
    cursor = conn.execute("SELECT id FROM phrases WHERE en_original = ?", (en_text,))
    if cursor.fetchone():
        conn.close()
        return jsonify({'error': 'Phrase already exists'}), 409
    
    new_id = str(uuid.uuid4())
    fields = ['id', 'en_original']
    values = [new_id, en_text]
    placeholders = ['?', '?']
    
    for lang_code in AVAILABLE_LANGUAGES:
        field = f'{lang_code}_original'
        if field in data and data[field].strip():
            fields.append(field)
            values.append(data[field].strip())
            placeholders.append('?')
    
    query = f"INSERT INTO phrases ({', '.join(fields)}) VALUES ({', '.join(placeholders)})"
    
    try:
        conn.execute(query, values)
        conn.commit()
    except Exception as e:
        return jsonify({'error': f'Database error: {str(e)}'}), 500
    finally:
        conn.close()
    
    # Clear phrase cache after adding new phrase
    try:
        from sds_translator_v4 import SDSTranslator
        SDSTranslator.clear_cache()
    except Exception as e:
        print(f"Could not clear phrase cache: {e}")
    
    return jsonify({'success': True, 'id': new_id, 'message': 'Phrase added'})

@database_bp.route('/api/phrases/bulk/upload', methods=['POST'])
def bulk_upload_phrases_route():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    file = request.files['file']
    if not file.filename.endswith('.txt'):
        return jsonify({'error': 'Only .txt files allowed'}), 400
    source_lang = request.form.get('source_lang', 'en')
    text = file.read().decode('utf-8')
    return bulk_process(text, source_lang)


def bulk_process(text, source_lang):
    from utils import parse_flag_format, LANG_TO_COLUMN
    phrases, detected_source = parse_flag_format(text)
    if source_lang == 'auto':
        source_lang = detected_source
    if not phrases:
        return jsonify({'error': 'No valid phrases found in text'}), 400

    source_text = next((t for lang, t in phrases if lang == source_lang), None)
    if not source_text:
        return jsonify({'error': f'No source phrase found for language: {source_lang}'}), 400

    translations = {lang: t for lang, t in phrases if lang != source_lang}
    source_column = LANG_TO_COLUMN.get(source_lang, 'en_original')

    conn = get_db_connection()
    updated_count = created_count = 0
    try:
        existing = conn.execute(f'SELECT id FROM phrases WHERE {source_column} = ?', (source_text,)).fetchone()
        if existing:
            update_fields = [f"{LANG_TO_COLUMN[lang]} = ?" for lang in translations if lang in LANG_TO_COLUMN]
            values = [translations[lang] for lang in translations if lang in LANG_TO_COLUMN]
            if update_fields:
                conn.execute(f"UPDATE phrases SET {', '.join(update_fields)} WHERE id = ?", values + [existing['id']])
                updated_count = 1
        else:
            import uuid
            new_id = str(uuid.uuid4())
            fields = ['id', source_column] + [LANG_TO_COLUMN[lang] for lang in translations if lang in LANG_TO_COLUMN]
            values = [new_id, source_text] + [translations[lang] for lang in translations if lang in LANG_TO_COLUMN]
            conn.execute(f"INSERT INTO phrases ({', '.join(fields)}) VALUES ({', '.join(['?']*len(fields))})", values)
            created_count = 1
        conn.commit()
    except Exception as e:
        return jsonify({'error': f'Database error: {str(e)}'}), 500
    finally:
        conn.close()

    return jsonify({'success': True, 'updated': updated_count, 'created': created_count,
                    'source_language': source_lang, 'source_phrase': source_text[:100]})


@database_bp.route('/api/cache/clear', methods=['POST'])
def clear_cache_route():
    """Manually clear the phrase cache."""
    try:
        from sds_translator_v4 import SDSTranslator
        SDSTranslator.clear_cache()
        return jsonify({'success': True, 'message': 'Phrase cache cleared'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@database_bp.route('/api/phrases/bulk/update', methods=['POST'])
def bulk_update_phrases_route():
    data = request.json
    text = data.get('text', '').strip()
    source_lang = data.get('source_lang', 'en')
    
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    
    phrases, detected_source = parse_flag_format(text)
    if source_lang == 'auto':
        source_lang = detected_source

    if not phrases:
        return jsonify({'error': 'No valid phrases found in text'}), 400
    
    conn = get_db_connection()
    updated_count = 0
    created_count = 0
    
    source_text = None
    translations = {}
    
    for lang, phrase_text in phrases:
        if lang == source_lang:
            source_text = phrase_text
        else:
            translations[lang] = phrase_text
    
    if not source_text:
        return jsonify({'error': f'No source phrase found for language: {source_lang}'}), 400
    
    source_column = LANG_TO_COLUMN.get(source_lang, 'en_original')
    
    try:
        cursor = conn.execute(f"SELECT id FROM phrases WHERE {source_column} = ?", (source_text,))
        existing = cursor.fetchone()
        
        if existing:
            phrase_id = existing['id']
            update_fields = [f"{LANG_TO_COLUMN[lang]} = ?" for lang in translations]
            values = list(translations.values())
            values.append(phrase_id)
            if update_fields:
                conn.execute(f"UPDATE phrases SET {', '.join(update_fields)} WHERE id = ?", values)
                updated_count += 1
        else:
            new_id = str(uuid.uuid4())
            fields = ['id', source_column] + [LANG_TO_COLUMN[lang] for lang in translations]
            values = [new_id, source_text] + list(translations.values())
            placeholders = ', '.join(['?'] * len(fields))
            conn.execute(f"INSERT INTO phrases ({', '.join(fields)}) VALUES ({placeholders})", values)
            created_count += 1
        
        conn.commit()
    except Exception as e:
        return jsonify({'error': f'Database error: {str(e)}'}), 500
    finally:
        conn.close()
        
    # Clear phrase cache after database update
    try:
        from sds_translator_v4 import SDSTranslator
        SDSTranslator.clear_cache()
        print("Phrase cache cleared after bulk update")
    except Exception as e:
        print(f"Could not clear phrase cache: {e}")
    
    return jsonify({
        'success': True,
        'updated': updated_count,
        'created': created_count,
        'source_language': source_lang,
        'source_phrase': source_text[:100],
        'cache_cleared': True
    })

@database_bp.route('/api/stats', methods=['GET'])
def get_stats_route():
    """Get database statistics."""
    conn = get_db_connection()
    
    try:
        cursor = conn.execute("SELECT COUNT(*) as count FROM phrases")
        total_phrases = cursor.fetchone()['count']
        
        lang_stats = {}
        for lang_code in AVAILABLE_LANGUAGES:
            column = LANG_TO_COLUMN.get(lang_code)
            if column:
                cursor = conn.execute(f"SELECT COUNT(*) as count FROM phrases WHERE {column} IS NOT NULL AND {column} != ''")
                lang_stats[lang_code] = cursor.fetchone()['count']
        
        conn.close()
        
        return jsonify({
            'total_phrases': total_phrases,
            'per_language': lang_stats
        })
        
    except Exception as e:
        conn.close()
        return jsonify({'error': str(e)}), 500


def get_phrases_for_export(phrase_ids=None, lang='de'):
    """Helper function to get phrases for export."""
    lang_col = f"{lang}_original"
    conn = get_db_connection()
    
    try:
        if phrase_ids and len(phrase_ids) > 0:
            placeholders = ', '.join(['?'] * len(phrase_ids))
            cursor = conn.execute(
                f"""SELECT id, en_original, {lang_col} as translation 
                    FROM phrases WHERE id IN ({placeholders}) ORDER BY en_original""",
                phrase_ids
            )
        else:
            cursor = conn.execute(
                f"""SELECT id, en_original, {lang_col} as translation 
                    FROM phrases ORDER BY en_original"""
            )
        
        phrases = [dict(row) for row in cursor.fetchall()]
        conn.close()
        return phrases
    except Exception as e:
        conn.close()
        raise e


@database_bp.route('/api/phrases/export', methods=['POST'])
def export_phrases_route():
    """Export phrases to PDF, HTML, or DOCX format."""
    data = request.json or {}
    phrase_ids = data.get('ids', [])
    export_format = data.get('format', 'html')
    lang = data.get('lang', 'de')
    
    if export_format not in ['pdf', 'html', 'docx']:
        return jsonify({'error': 'Invalid export format. Use pdf, html, or docx'}), 400
    
    try:
        phrases = get_phrases_for_export(phrase_ids if phrase_ids else None, lang)
        
        if not phrases:
            return jsonify({'error': 'Keine Phrasen zum Exportieren vorhanden'}), 400
        
        if export_format == 'html':
            return export_as_html(phrases, lang)
        elif export_format == 'pdf':
            return export_as_pdf(phrases, lang)
        elif export_format == 'docx':
            return export_as_docx(phrases, lang)
            
    except Exception as e:
        return jsonify({'error': f'Export error: {str(e)}'}), 500


def export_as_html(phrases, lang):
    """Export phrases as HTML."""
    from datetime import datetime
    
    lang_name = {'de': 'Deutsch', 'en': 'English', 'fr': 'Français', 'es': 'Español', 'it': 'Italiano', 'pt': 'Português'}.get(lang, lang.upper())
    
    html_content = f"""<!DOCTYPE html>
<html lang="de">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Übersetzungen Export - {datetime.now().strftime('%d.%m.%Y')}</title>
    <style>
        * {{ box-sizing: border-box; margin: 0; padding: 0; }}
        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; line-height: 1.6; color: #333; background: #f5f5f5; padding: 20px; }}
        .container {{ max-width: 1200px; margin: 0 auto; background: white; padding: 40px; border-radius: 8px; box-shadow: 0 2px 10px rgba(0,0,0,0.1); }}
        h1 {{ color: #76B82A; margin-bottom: 10px; font-size: 28px; }}
        .meta {{ color: #666; font-size: 14px; margin-bottom: 30px; padding-bottom: 20px; border-bottom: 2px solid #76B82A; }}
        .stats {{ display: flex; gap: 20px; margin-bottom: 30px; }}
        .stat-box {{ background: #f0f0f0; padding: 15px 20px; border-radius: 6px; flex: 1; }}
        .stat-box strong {{ display: block; font-size: 24px; color: #76B82A; }}
        .stat-box span {{ font-size: 12px; color: #666; text-transform: uppercase; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 20px; }}
        th {{ background: #76B82A; color: white; padding: 12px; text-align: left; font-weight: 600; }}
        td {{ padding: 12px; border-bottom: 1px solid #eee; }}
        tr:hover {{ background: #f9f9f9; }}
        .id-cell {{ font-family: monospace; color: #999; font-size: 12px; }}
        .en-text {{ font-weight: 500; color: #333; }}
        .translation-text {{ color: #76B82A; font-weight: 500; }}
        @media (max-width: 768px) {{ .container {{ padding: 20px; }} table {{ font-size: 14px; }} th, td {{ padding: 8px; }} .stats {{ flex-direction: column; }} }}
    </style>
</head>
<body>
    <div class="container">
        <h1>📚 Übersetzungen Export</h1>
        <div class="meta">
            <strong>Exportdatum:</strong> {datetime.now().strftime('%d.%m.%Y um %H:%M Uhr')}<br>
            <strong>Zielsprache:</strong> {lang_name}
        </div>
        <div class="stats">
            <div class="stat-box"><strong>{len(phrases)}</strong><span>Exportierte Phrasen</span></div>
        </div>
        <table>
            <thead><tr><th style="width: 15%;">ID</th><th style="width: 42%;">Englisch (Original)</th><th style="width: 43%;">{lang_name}</th></tr></thead>
            <tbody>
"""
    for phrase in phrases:
        en_text = phrase.get('en_original', '')
        translation = phrase.get('translation', '')
        phrase_id = phrase.get('id', '')[:8]
        html_content += f'<tr><td class="id-cell">{phrase_id}...</td><td class="en-text">{en_text}</td><td class="translation-text">{translation}</td></tr>\n'
    
    html_content += """            </tbody>
        </table>
    </div>
</body>
</html>"""
    
    response = make_response(html_content)
    response.headers['Content-Type'] = 'text/html; charset=utf-8'
    response.headers['Content-Disposition'] = f"attachment; filename=translations_export_{datetime.now().strftime('%Y%m%d')}.html"
    return response


def export_as_pdf(phrases, lang):
    """Export phrases as PDF using WeasyPrint."""
    from datetime import datetime
    from weasyprint import HTML
    
    lang_name = {'de': 'Deutsch', 'en': 'English', 'fr': 'Français', 'es': 'Español', 'it': 'Italiano', 'pt': 'Português'}.get(lang, lang.upper())
    
    html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <style>
        @page {{ size: A4; margin: 2cm; }}
        body {{ font-family: Arial, sans-serif; font-size: 10pt; line-height: 1.4; }}
        .header {{ text-align: center; margin-bottom: 20px; padding-bottom: 15px; border-bottom: 2px solid #76B82A; }}
        h1 {{ color: #76B82A; font-size: 24pt; margin-bottom: 5px; }}
        .meta {{ color: #666; font-size: 9pt; }}
        .stats {{ background: #f5f5f5; padding: 10px; margin-bottom: 20px; border-radius: 4px; }}
        table {{ width: 100%; border-collapse: collapse; margin-top: 15px; }}
        th {{ background: #76B82A; color: white; padding: 8px; text-align: left; font-weight: bold; }}
        td {{ padding: 6px 8px; border-bottom: 1px solid #ddd; }}
        tr:nth-child(even) {{ background: #f9f9f9; }}
        .id-col {{ font-size: 8pt; color: #999; width: 12%; }}
        .en-col {{ width: 44%; }}
        .lang-col {{ width: 44%; color: #76B82A; }}
    </style>
</head>
<body>
    <div class="header">
        <h1>Übersetzungen Export</h1>
        <div class="meta"><p><strong>Exportdatum:</strong> {datetime.now().strftime('%d.%m.%Y')}</p><p><strong>Zielsprache:</strong> {lang_name}</p></div>
    </div>
    <div class="stats"><strong>Anzahl exportierter Phrasen:</strong> {len(phrases)}</div>
    <table>
        <thead><tr><th class="id-col">ID</th><th class="en-col">Englisch (Original)</th><th class="lang-col">{lang_name}</th></tr></thead>
        <tbody>
"""
    for phrase in phrases:
        en_text = phrase.get('en_original', '')
        translation = phrase.get('translation', '')
        phrase_id = phrase.get('id', '')[:8]
        html_content += f'<tr><td class="id-col">{phrase_id}...</td><td>{en_text}</td><td class="lang-col">{translation}</td></tr>\n'
    
    html_content += """        </tbody>
    </table>
</body>
</html>"""
    
    pdf_file = HTML(string=html_content).write_pdf()
    
    response = make_response(pdf_file)
    response.headers['Content-Type'] = 'application/pdf'
    response.headers['Content-Disposition'] = f"attachment; filename=translations_export_{datetime.now().strftime('%Y%m%d')}.pdf"
    return response


def export_as_docx(phrases, lang):
    """Export phrases as DOCX using python-docx."""
    from datetime import datetime
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    from io import BytesIO
    
    lang_name = {'de': 'Deutsch', 'en': 'English', 'fr': 'Français', 'es': 'Español', 'it': 'Italiano', 'pt': 'Português'}.get(lang, lang.upper())
    
    doc = Document()
    
    # Title
    title = doc.add_heading('Übersetzungen Export', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(118, 184, 42)
    
    # Meta
    doc.add_paragraph(f'Exportdatum: {datetime.now().strftime("%d.%m.%Y")}')
    doc.add_paragraph(f'Zielsprache: {lang_name}')
    doc.add_paragraph(f'Anzahl Phrasen: {len(phrases)}')
    doc.add_paragraph()
    
    # Table
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = 'ID'
    header_cells[1].text = 'Englisch (Original)'
    header_cells[2].text = lang_name
    
    for header_cell in header_cells:
        header_cell.paragraphs[0].runs[0].font.bold = True
        header_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(118, 184, 42)
        header_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading_elm = parse_xml(r'<w:shd {} w:fill="76B82A"/>'.format(nsdecls('w')))
        header_cell._tc.get_or_add_tcPr().append(shading_elm)
        for paragraph in header_cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
    
    # Data rows
    for phrase in phrases:
        en_text = phrase.get('en_original', '')
        translation = phrase.get('translation', '')
        phrase_id = phrase.get('id', '')[:8] + '...'
        
        row_cells = table.add_row().cells
        row_cells[0].text = phrase_id
        row_cells[1].text = en_text
        row_cells[2].text = translation
        
        row_cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        row_cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(128, 128, 128)
        row_cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(118, 184, 42)
    
    # Save
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    response = make_response(doc_buffer.getvalue())
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    response.headers['Content-Disposition'] = f"attachment; filename=translations_export_{datetime.now().strftime('%Y%m%d')}.docx"
    return response
 
